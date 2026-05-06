from pathlib import Path
from typing import Any

from docx import Document
from docx.shared import Inches


def stringify(value: Any) -> str:
    if value is None:
        return ""

    if isinstance(value, bool):
        return "Yes" if value else "No"

    if isinstance(value, list):
        return ", ".join(str(item) for item in value)

    if isinstance(value, dict):
        return ", ".join(f"{key}: {val}" for key, val in value.items())

    return str(value)


def add_key_value_table(document: Document, data: dict[str, Any]) -> None:
    table = document.add_table(rows=1, cols=2)
    table.style = "Table Grid"

    header_cells = table.rows[0].cells
    header_cells[0].text = "Field"
    header_cells[1].text = "Value"

    for key, value in data.items():
        row_cells = table.add_row().cells
        row_cells[0].text = stringify(key)
        row_cells[1].text = stringify(value)


def add_bullet_list(document: Document, items: list[Any]) -> None:
    for item in items:
        document.add_paragraph(stringify(item), style="List Bullet")


def add_attack_flow_table(document: Document, attack_flow: list[dict[str, Any]]) -> None:
    table = document.add_table(rows=1, cols=2)
    table.style = "Table Grid"

    header_cells = table.rows[0].cells
    header_cells[0].text = "Step"
    header_cells[1].text = "Description"

    for item in attack_flow:
        row_cells = table.add_row().cells
        row_cells[0].text = stringify(item.get("step", ""))
        row_cells[1].text = stringify(item.get("description", ""))


def add_attack_mapping_table(document: Document, mappings: list[dict[str, Any]]) -> None:
    table = document.add_table(rows=1, cols=4)
    table.style = "Table Grid"

    header_cells = table.rows[0].cells
    header_cells[0].text = "Tactic"
    header_cells[1].text = "Technique ID"
    header_cells[2].text = "Technique Name"
    header_cells[3].text = "Evidence"

    for item in mappings:
        row_cells = table.add_row().cells
        row_cells[0].text = stringify(item.get("tactic", ""))
        row_cells[1].text = stringify(item.get("technique_id", ""))
        row_cells[2].text = stringify(item.get("technique_name", ""))
        row_cells[3].text = stringify(item.get("evidence", ""))


def add_rules_table(document: Document, rules: dict[str, Any]) -> None:
    for rule_type, rule_items in rules.items():
        document.add_heading(rule_type.upper(), level=3)

        if not rule_items:
            document.add_paragraph(f"No {rule_type} rules provided.")
            continue

        table = document.add_table(rows=1, cols=3)
        table.style = "Table Grid"

        header_cells = table.rows[0].cells
        header_cells[0].text = "Name"
        header_cells[1].text = "Path"
        header_cells[2].text = "Description"

        for rule in rule_items:
            row_cells = table.add_row().cells
            row_cells[0].text = stringify(rule.get("name", ""))
            row_cells[1].text = stringify(rule.get("path", ""))
            row_cells[2].text = stringify(rule.get("description", ""))


def render_docx_report(data: dict[str, Any], output_path: Path) -> Path:
    document = Document()

    section = document.sections[0]
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.8)
    section.left_margin = Inches(0.8)
    section.right_margin = Inches(0.8)

    document.add_heading(data.get("title", "Security Report"), level=0)

    meta = {
        "Author": data.get("author", ""),
        "Date": data.get("date", ""),
        "Report Type": str(data.get("report_type", "")).upper(),
    }
    add_key_value_table(document, meta)

    document.add_heading("1. Executive Summary", level=1)
    summary = data.get("summary", {}) or {}
    document.add_paragraph(summary.get("overview", "No executive summary provided."))

    if summary.get("key_risk"):
        paragraph = document.add_paragraph()
        paragraph.add_run("Key Risk: ").bold = True
        paragraph.add_run(stringify(summary.get("key_risk")))

    document.add_heading("2. Target Information", level=1)
    target = data.get("target", {}) or {}
    if target:
        add_key_value_table(document, target)
    else:
        document.add_paragraph("No target information provided.")

    document.add_heading("3. Severity Assessment", level=1)
    severity = data.get("severity", {}) or {}
    if severity:
        add_key_value_table(document, severity)
    else:
        document.add_paragraph("No severity information provided.")

    document.add_heading("4. Technical Analysis", level=1)
    document.add_paragraph(
        summary.get(
            "technical_details",
            "Technical details should describe the vulnerable component, root cause, attack condition, and security impact.",
        )
    )

    document.add_heading("5. Attack Flow", level=1)
    attack_flow = data.get("attack_flow", []) or []
    if attack_flow:
        add_attack_flow_table(document, attack_flow)
    else:
        document.add_paragraph("No attack flow provided.")

    document.add_heading("6. Detection Points", level=1)
    detection = data.get("detection", {}) or {}
    if detection:
        for category, items in detection.items():
            document.add_heading(str(category).title(), level=2)
            if isinstance(items, list):
                add_bullet_list(document, items)
            else:
                document.add_paragraph(stringify(items))
    else:
        document.add_paragraph("No detection points provided.")

    document.add_heading("7. Indicators of Compromise", level=1)
    iocs = data.get("iocs", {}) or {}
    if iocs:
        for category, items in iocs.items():
            document.add_heading(str(category).title(), level=2)
            if isinstance(items, list) and items:
                add_bullet_list(document, items)
            else:
                document.add_paragraph("None")
    else:
        document.add_paragraph("No IOC provided.")

    document.add_heading("8. MITRE ATT&CK Mapping", level=1)
    attack_mapping = data.get("attack_mapping", []) or []
    if attack_mapping:
        add_attack_mapping_table(document, attack_mapping)
    else:
        document.add_paragraph("No ATT&CK mapping provided.")

    document.add_heading("9. Detection Rules", level=1)
    rules = data.get("rules", {}) or {}
    if rules:
        add_rules_table(document, rules)
    else:
        document.add_paragraph("No detection rules provided.")

    document.add_heading("10. Mitigation and Remediation", level=1)
    mitigation = data.get("mitigation", []) or []
    if mitigation:
        add_bullet_list(document, mitigation)
    else:
        document.add_paragraph("No mitigation guidance provided.")

    document.add_heading("11. References", level=1)
    references = data.get("references", []) or []
    if references:
        add_bullet_list(document, references)
    else:
        document.add_paragraph("No references provided.")

    document.save(output_path)
    return output_path